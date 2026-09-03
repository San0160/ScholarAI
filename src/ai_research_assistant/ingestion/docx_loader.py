from pathlib import Path

from docx import Document as DocxDocument

from ai_research_assistant.entity.document import Document
from ai_research_assistant.ingestion.base_loader import BaseLoader


class DocxLoader(BaseLoader):

    def load(self, file_path: str) -> list[Document]:

        file_path = Path(file_path)

        doc = DocxDocument(file_path)

        documents = []

        current_section = None
        buffer = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            style_name = paragraph.style.name.lower()

            is_heading = style_name.startswith(
                "heading"
            )

            if is_heading:

                if buffer:

                    documents.append(self._create_document("\n".join(buffer), file_path, current_section))

                    buffer = []

                current_section = text.lower()

                continue

            buffer.append(text)

        if buffer:

            documents.append(self._create_document("\n".join(buffer), file_path, current_section))

        return documents

    def _create_document(self, text: str, file_path: Path, section: str | None) -> Document:

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": "docx"
        }

        if section:
            metadata["section"] = section

        return Document(
            page_content=text,
            metadata=metadata
        )